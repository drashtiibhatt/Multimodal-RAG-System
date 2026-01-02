"""
DOCX Parser - Extracts text from Microsoft Word documents.

Supports:
- Paragraph extraction
- Table extraction and formatting
- Document metadata
- Structured content preservation
"""

from pathlib import Path
from typing import List, Dict, Any
import logging

try:
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base import BaseParser, Document

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for Microsoft Word (.docx) documents."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self, extract_tables: bool = True):
        """
        Initialize DOCX parser.

        Args:
            extract_tables: Whether to extract tables from document
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not available. Install with: pip install python-docx"
            )

        self.extract_tables = extract_tables
        logger.info(f"DOCXParser initialized (extract_tables={extract_tables})")

    def parse(self, file_path: str) -> List[Document]:
        """
        Parse DOCX file and extract text content.

        Args:
            file_path: Path to DOCX file

        Returns:
            List containing a single Document with all content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a DOCX
        """
        path = self._validate_file_exists(file_path)

        # Validate extension
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"File is not a DOCX: {file_path}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        try:
            # Load document
            doc = DocxDocument(str(path))

            # Extract content
            content_parts = []

            # Extract paragraphs
            paragraphs = self._extract_paragraphs(doc)
            if paragraphs:
                content_parts.append("=== DOCUMENT CONTENT ===\n")
                content_parts.append(paragraphs)

            # Extract tables if enabled
            if self.extract_tables:
                tables = self._extract_tables(doc)
                if tables:
                    content_parts.append("\n\n=== TABLES ===\n")
                    content_parts.append(tables)

            # Combine all content
            full_content = "\n".join(content_parts)

            # Extract metadata
            metadata = self._extract_metadata(path, doc)

            # Create document
            document = Document(
                page_content=full_content,
                metadata=metadata
            )

            logger.info(
                f"Extracted {len(full_content)} characters from {path.name} "
                f"({metadata.get('paragraph_count', 0)} paragraphs, "
                f"{metadata.get('table_count', 0)} tables)"
            )

            return [document]

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {str(e)}")
            raise

    def _extract_paragraphs(self, doc: DocxDocument) -> str:
        """
        Extract all paragraphs from document.

        Args:
            doc: python-docx Document object

        Returns:
            Formatted paragraph text
        """
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()

            # Skip empty paragraphs
            if not text:
                continue

            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                # Add extra spacing for headings
                paragraphs.append(f"\n## {text}\n")
            else:
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def _extract_tables(self, doc: DocxDocument) -> str:
        """
        Extract and format all tables from document.

        Args:
            doc: python-docx Document object

        Returns:
            Formatted table text
        """
        if not doc.tables:
            return ""

        table_texts = []

        for table_idx, table in enumerate(doc.tables, start=1):
            table_text = self._format_table(table, table_idx)
            if table_text:
                table_texts.append(table_text)

        return "\n\n".join(table_texts)

    def _format_table(self, table: Table, table_number: int) -> str:
        """
        Format a single table as text.

        Args:
            table: python-docx Table object
            table_number: Table index in document

        Returns:
            Formatted table text
        """
        lines = [f"Table {table_number}:"]

        for row_idx, row in enumerate(table.rows):
            # Extract cell values
            cells = [cell.text.strip() for cell in row.cells]

            # Skip empty rows
            if not any(cells):
                continue

            # Format as pipe-separated values
            if row_idx == 0:
                # Header row
                lines.append("| " + " | ".join(cells) + " |")
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")
            else:
                # Data row
                lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)

    def _extract_metadata(
        self,
        path: Path,
        doc: DocxDocument
    ) -> Dict[str, Any]:
        """
        Extract metadata from DOCX document.

        Args:
            path: Path to DOCX file
            doc: python-docx Document object

        Returns:
            Dictionary with document metadata
        """
        # Get core properties
        core_props = doc.core_properties

        metadata = {
            "source": str(path),
            "file_name": path.name,
            "file_type": "docx",
            "size_bytes": path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        # Add core properties if available
        try:
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
        except Exception as e:
            logger.warning(f"Failed to extract some core properties: {str(e)}")

        return metadata

    @classmethod
    def supports_file(cls, file_path: str) -> bool:
        """
        Check if this parser supports the given file.

        Args:
            file_path: Path to file

        Returns:
            True if file extension is supported
        """
        path = Path(file_path)
        return path.suffix.lower() in cls.SUPPORTED_EXTENSIONS
